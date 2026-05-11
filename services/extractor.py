"""
Text Extraction Service
Handles PDF (Google Vision API), DOCX (python-docx + LibreOffice .doc→.docx),
and Images (Google Vision API).
"""

import base64
import io
import logging
import os
import subprocess
import tempfile
import time
import shutil
from typing import Literal

logger = logging.getLogger(__name__)





# ── Google Vision Client ──────────────────────────────────────────────────────

def _get_vision_client():
    """
    Returns an authenticated Google Vision client.
    """
    try:
        from google.cloud import vision
        from google.oauth2 import service_account
    except ImportError:
        raise ImportError("[ERROR] google-cloud-vision not installed")

    raw_json = os.environ.get("GOOGLE_CREDENTIALS_JSON")

    if raw_json:
        import json
        try:
            info = json.loads(raw_json)
            credentials = service_account.Credentials.from_service_account_info(
                info,
                scopes=["https://www.googleapis.com/auth/cloud-platform"],
            )
            print("[INFO] Google Vision: using JSON env credentials")
            return vision.ImageAnnotatorClient(credentials=credentials)
        except Exception as e:
            print(f"[ERROR] Invalid GOOGLE_CREDENTIALS_JSON: {e}")
            raise

    print("[WARN] Google Vision: using default credentials (ADC)")
    return vision.ImageAnnotatorClient()


# ── PDF Extraction ────────────────────────────────────────────────────────────

def _extract_pdf(data: bytes) -> tuple[str, int]:
    try:
        print("PDF -> Using Google Vision OCR")
        return _extract_pdf_vision(data)
    except Exception as e:
        logger.warning(f"Google Vision PDF extraction failed ({e}), falling back to PyMuPDF")
        print("[WARN] PDF -> Using PyMuPDF fallback")
        return _extract_pdf_pymupdf(data), 0


def _extract_pdf_vision(data: bytes) -> tuple[str, int]:
    import fitz
    from google.cloud import vision

    client = _get_vision_client()
    doc = fitz.open(stream=data, filetype="pdf")
    pages_text = []
    vision_units = 0

    for page_num, page in enumerate(doc, start=1):
        mat = fitz.Matrix(150 / 72, 150 / 72)
        pix = page.get_pixmap(matrix=mat)
        png_bytes = pix.tobytes("png")

        image = vision.Image(content=png_bytes)
        response = client.document_text_detection(image=image)
        vision_units += 1

        if response.error.message:
            raise RuntimeError(f"Vision API error on page {page_num}: {response.error.message}")

        text = response.full_text_annotation.text
        if text.strip():
            pages_text.append(f"--- Page {page_num} ---\n{text.strip()}")

    doc.close()
    return "\n\n".join(pages_text), vision_units


def _extract_pdf_pymupdf(data: bytes) -> str:
    import fitz

    doc = fitz.open(stream=data, filetype="pdf")
    pages = []
    for page_num, page in enumerate(doc, start=1):
        text = page.get_text("text")
        if text.strip():
            pages.append(f"--- Page {page_num} ---\n{text}")
    doc.close()
    return "\n\n".join(pages)


# ── DOCX Extraction ───────────────────────────────────────────────────────────

def _convert_doc_to_docx(doc_bytes: bytes) -> bytes:
    with tempfile.TemporaryDirectory() as tmpdir:
        src = os.path.join(tmpdir, "input.doc")
        with open(src, "wb") as f:
            f.write(doc_bytes)

        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "docx", "--outdir", tmpdir, src],
            capture_output=True,
            text=True,
            timeout=60,
        )

        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")

        out = src.replace(".doc", ".docx")
        with open(out, "rb") as f:
            return f.read()


def _extract_docx(data: bytes, file_name: str) -> str:
    # pyrefly: ignore [missing-import]
    from docx import Document

    if file_name.lower().endswith(".doc") and not file_name.lower().endswith(".docx"):
        logger.info("Legacy .doc detected – converting via LibreOffice")
        data = _convert_doc_to_docx(data)

    doc = Document(io.BytesIO(data))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(f"[TABLE ROW] {row_text}")

    return "\n".join(paragraphs)


# ── Image OCR ─────────────────────────────────────────────────────────────────

def _extract_image_vision(data: bytes) -> tuple[str, int]:
    from google.cloud import vision

    client = _get_vision_client()
    image = vision.Image(content=data)
    response = client.document_text_detection(image=image)

    if response.error.message:
        raise RuntimeError(f"Vision API error: {response.error.message}")

    text = response.full_text_annotation.text
    logger.info(f"Google Vision extracted {len(text)} chars")

    print("[INFO] OCR Engine: Google Vision")

    return text.strip(), 1





def _extract_image(data: bytes) -> tuple[str, int]:
    for attempt in range(2):
        try:
            print(f"Trying Google Vision (Attempt {attempt+1})")

            text, units = _extract_image_vision(data)

            if len(text) < 10:
                raise ValueError("Too little text")

            return text, units

        except Exception as e:
            logger.warning(f"Vision attempt {attempt+1} failed: {e}")
            time.sleep(1)

    raise RuntimeError("All Google Vision attempts failed. No fallback available.")


# ── Public Interface ──────────────────────────────────────────────────────────

def extract_text(
    file_name: str,
    file_type: Literal["pdf", "docx", "image"],
    file_base64: str,
) -> tuple[str, int]:

    if "," in file_base64:
        file_base64 = file_base64.split(",", 1)[1]

    data = base64.b64decode(file_base64)

    logger.info(
        f"Extracting text | file={file_name} | type={file_type} | size={len(data)} bytes"
    )

    if file_type == "pdf":
        return _extract_pdf(data)
    elif file_type == "docx":
        return _extract_docx(data, file_name), 0
    elif file_type == "image":
        return _extract_image(data)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
    
